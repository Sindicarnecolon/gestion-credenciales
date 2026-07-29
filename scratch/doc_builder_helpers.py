import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_header_footer(doc, title):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = f"Guía de Desarrollo - {title}"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.size = Pt(8.5)
    hp.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Aplicación de Credencial Digital de Afiliados | Documento Confidencial"
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.runs[0].font.size = Pt(8.5)
    fp.runs[0].font.color.rgb = RGBColor(128, 128, 128)

def format_paragraph(p, space_before=0, space_after=6, line_spacing=1.15):
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.keep_with_next = True
    if level == 1:
        format_paragraph(h, space_before=16, space_after=8)
        for r in h.runs:
            r.font.size = Pt(18)
            r.font.bold = True
            r.font.color.rgb = RGBColor(13, 71, 161) # Primary Blue
    elif level == 2:
        format_paragraph(h, space_before=12, space_after=6)
        for r in h.runs:
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0, 137, 123) # Teal Accent
    elif level == 3:
        format_paragraph(h, space_before=8, space_after=4)
        for r in h.runs:
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = RGBColor(60, 60, 60)
    return h

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    p = cell.paragraphs[0]
    format_paragraph(p, space_before=2, space_after=2, line_spacing=1.0)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(33, 33, 33)
    doc.add_paragraph() # spacing

def add_callout(doc, title, text, bg_hex="E3F2FD", border_hex="2196F3"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border styling
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="36" w:space="0" w:color="{border_hex}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    format_paragraph(p, space_before=2, space_after=4)
    r_title = p.add_run(f"💡 {title}\n")
    r_title.bold = True
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor(13, 71, 161)
    
    r_text = p.add_run(text)
    r_text.font.size = Pt(10)
    r_text.font.color.rgb = RGBColor(40, 40, 40)
    doc.add_paragraph()

print("Helper script ready.")
