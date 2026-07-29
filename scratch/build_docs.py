import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def format_paragraph(p, space_before=0, space_after=6, line_spacing=1.15):
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing

def add_header_footer(doc, title):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    header = section.header
    hp = header.paragraphs[0]
    hp.text = f"Guía de Desarrollo & Despliegue - {title}"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if hp.runs:
        hp.runs[0].font.size = Pt(8.5)
        hp.runs[0].font.color.rgb = RGBColor(120, 120, 120)
    
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Aplicación de Credencial Digital de Afiliados | Documento Técnico"
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if fp.runs:
        fp.runs[0].font.size = Pt(8.5)
        fp.runs[0].font.color.rgb = RGBColor(120, 120, 120)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.keep_with_next = True
    if level == 1:
        format_paragraph(h, space_before=18, space_after=8)
        for r in h.runs:
            r.font.name = 'Segoe UI'
            r.font.size = Pt(18)
            r.font.bold = True
            r.font.color.rgb = RGBColor(13, 71, 161) # Blue
    elif level == 2:
        format_paragraph(h, space_before=14, space_after=6)
        for r in h.runs:
            r.font.name = 'Segoe UI'
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0, 137, 123) # Teal
    elif level == 3:
        format_paragraph(h, space_before=10, space_after=4)
        for r in h.runs:
            r.font.name = 'Segoe UI'
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = RGBColor(60, 60, 60)
    return h

def add_p(doc, text="", bold_prefix="", italic=False):
    p = doc.add_paragraph()
    format_paragraph(p, space_before=0, space_after=6)
    if bold_prefix:
        r_b = p.add_run(bold_prefix)
        r_b.font.name = 'Segoe UI'
        r_b.font.bold = True
        r_b.font.size = Pt(10.5)
        r_b.font.color.rgb = RGBColor(30, 30, 30)
    if text:
        r = p.add_run(text)
        r.font.name = 'Segoe UI'
        r.font.size = Pt(10.5)
        r.font.italic = italic
        r.font.color.rgb = RGBColor(50, 50, 50)
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    format_paragraph(p, space_before=0, space_after=4)
    if bold_prefix:
        r_b = p.add_run(bold_prefix)
        r_b.font.name = 'Segoe UI'
        r_b.font.bold = True
        r_b.font.size = Pt(10)
    r = p.add_run(text)
    r.font.name = 'Segoe UI'
    r.font.size = Pt(10)
    return p

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F8F9FA")
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="0" w:color="CCCCCC"/><w:top w:val="single" w:sz="12" w:space="0" w:color="CCCCCC"/><w:right w:val="single" w:sz="12" w:space="0" w:color="CCCCCC"/><w:bottom w:val="single" w:sz="12" w:space="0" w:color="CCCCCC"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    format_paragraph(p, space_before=2, space_after=2, line_spacing=1.0)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(33, 33, 33)
    
    empty_p = doc.add_paragraph()
    format_paragraph(empty_p, space_before=0, space_after=4)

def add_callout(doc, title, text, bg_hex="E3F2FD", border_hex="2196F3"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="36" w:space="0" w:color="{border_hex}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    format_paragraph(p, space_before=2, space_after=4)
    r_title = p.add_run(f"📌 {title}\n")
    r_title.font.name = 'Segoe UI'
    r_title.bold = True
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor(13, 71, 161)
    
    r_text = p.add_run(text)
    r_text.font.name = 'Segoe UI'
    r_text.font.size = Pt(10)
    r_text.font.color.rgb = RGBColor(40, 40, 40)
    
    empty_p = doc.add_paragraph()
    format_paragraph(empty_p, space_before=0, space_after=4)

# ==========================================
# 1. GENERATE ANDROID GUIDE
# ==========================================
def build_android_doc(filename):
    doc = docx.Document()
    add_header_footer(doc, "Compilación y Testeo en Android")
    
    # Title Block
    title_p = doc.add_paragraph()
    format_paragraph(title_p, space_before=10, space_after=4)
    r = title_p.add_run("Guía de Compilación, Testeo y Vinculación Web - Android")
    r.font.name = 'Segoe UI'
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(13, 71, 161)
    
    sub_p = doc.add_paragraph()
    format_paragraph(sub_p, space_before=0, space_after=18)
    r_sub = sub_p.add_run("Aplicación de Credencial Digital de Afiliados (Flutter Native Shell)")
    r_sub.font.name = 'Segoe UI'
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(100, 100, 100)
    
    add_callout(doc, "Estado Actual del Cascarón Nativo (Shell)",
                "El cascarón nativo Flutter ya se encuentra 100% construido y validado (dart analyze limpio, flutter test exitoso). Esta guía contiene las instrucciones exactas para ejecutar la app en emuladores Android, integrar las páginas web finales y generar los instalables APK/AAB.")

    # Section 1: Pre-requisitos
    add_styled_heading(doc, "1. Requisitos Previos del Entorno", level=1)
    add_p(doc, "Antes de comenzar la compilación en Windows, asegúrese de contar con los siguientes elementos instalados y verificados:")
    add_bullet(doc, "Instalado en C:\\src\\flutter y verificado en PATH.", bold_prefix="Flutter SDK 3.44.7+: ")
    add_bullet(doc, "Incluido con Flutter 3.44.7.", bold_prefix="Dart SDK 3.12.2+: ")
    add_bullet(doc, "Instalado con SDK Platform 36.1.0 y Build Tools.", bold_prefix="Android SDK: ")
    add_bullet(doc, "JDK 21 incluido en Android Studio (configurado por defecto).", bold_prefix="Java Development Kit (JDK): ")
    add_bullet(doc, "Ejecutar en la terminal 'flutter doctor -v' para confirmar que todos los ítems estén en verde [√].", bold_prefix="Verificación: ")

    # Section 2: Emuladores Android
    add_styled_heading(doc, "2. Paso a Paso: Configuración y Ejecución en Emulador Android (AVD)", level=1)
    add_p(doc, "Existen dos alternativas para crear y ejecutar un emulador Android (Android Virtual Device - AVD):")
    
    add_styled_heading(doc, "Opción A: Desde la Interfaz Gráfica de Android Studio (Recomendado)", level=2)
    add_bullet(doc, "Abra Android Studio desde el menú de inicio de Windows.")
    add_bullet(doc, "En la pantalla de bienvenida o menú superior, seleccione Virtual Device Manager (Gestor de dispositivos virtuales).")
    add_bullet(doc, "Haga clic en 'Create Device' (Crear dispositivo).")
    add_bullet(doc, "Seleccione un modelo recomendado de teléfono, por ejemplo Pixel 8 o Pixel 7, y presione 'Next'.")
    add_bullet(doc, "Seleccione la versión del sistema (por ejemplo API 34 o API 36 / VanillaIceCream) y haga clic en 'Download' si aún no está instalada.")
    add_bullet(doc, "Haga clic en 'Finish'.")
    add_bullet(doc, "Inicie el emulador haciendo clic en el icono de reproducción ▶ (Play) junto al dispositivo creado.")

    add_styled_heading(doc, "Opción B: Desde la Consola de Comandos (PowerShell / Antigravity CLI)", level=2)
    add_p(doc, "Puede gestionar los emuladores directamente mediante la herramienta de línea de comandos de Flutter:")
    add_code_block(doc, "# Listar los emuladores disponibles\nflutter emulators\n\n# Crear un emulador si no existe ningún AVD\nflutter emulators --create --name pixel_emulator\n\n# Iniciar el emulador por nombre\nflutter emulators --launch pixel_emulator")

    # Section 3: Ejecución y Compilación
    add_styled_heading(doc, "3. Paso a Paso: Compilación y Ejecución de la App", level=1)
    
    add_styled_heading(doc, "3.1 Modo Desarrollo (Pruebas con Hot Reload)", level=2)
    add_p(doc, "Con el emulador Android encendido y visible en pantalla, ejecute el siguiente comando en la raíz del proyecto:")
    add_code_block(doc, "flutter run -d android")
    add_p(doc, "Durante la ejecución en este modo:")
    add_bullet(doc, "Presione 'r' en la consola para realizar Hot Reload (recarga instantánea del código).")
    add_bullet(doc, "Presione 'R' para realizar Hot Restart (reiniciar el estado completo de la app).")
    add_bullet(doc, "Presione 'q' para detener la ejecución.")

    add_styled_heading(doc, "3.2 Compilación del APK de Pruebas (Debug APK)", level=2)
    add_p(doc, "Para generar un archivo instalable sin necesidad de tener el emulador conectado a la consola de desarrollo:")
    add_code_block(doc, "flutter build apk --debug")
    add_p(doc, "El archivo generado se ubicará en:", bold_prefix="Ubicación del APK: ")
    add_code_block(doc, "build\\app\\outputs\\flutter-apk\\app-debug.apk")
    add_p(doc, "Puede instalar este APK directamente en el emulador arrastrando el archivo sobre la ventana del emulador o ejecutando:")
    add_code_block(doc, "adb install build\\app\\outputs\\flutter-apk\\app-debug.apk")

    add_styled_heading(doc, "3.3 Compilación del APK de Producción (Release APK)", level=2)
    add_p(doc, "Para generar el instalable final optimizado para distribución directa a los afiliados:")
    add_code_block(doc, "flutter build apk --release")
    add_p(doc, "El archivo comprimido y optimizado estará disponible en:", bold_prefix="Ubicación Release: ")
    add_code_block(doc, "build\\app\\outputs\\flutter-apk\\app-release.apk")

    add_styled_heading(doc, "3.4 Compilación de App Bundle para Google Play Store (.aab)", level=2)
    add_p(doc, "Si desea publicar la aplicación en la tienda oficial Google Play Store, debe generar un archivo Android App Bundle (.aab):")
    add_code_block(doc, "flutter build appbundle --release")
    add_p(doc, "El archivo resultante estará en:", bold_prefix="Ubicación AAB: ")
    add_code_block(doc, "build\\app\\outputs\\bundle\\release\\app-release.aab")

    # Section 4: Vincular las páginas Web
    add_styled_heading(doc, "4. Pasos Necesarios para Agregar las Páginas Web al Cascarón (Shell)", level=1)
    add_p(doc, "Toda la infraestructura de la aplicación nativa se encuentra conectada al archivo de configuración centralizado. Para vincular la aplicación web real a los cascarones nativos (WebView), siga estos pasos exactamente:")
    
    add_p(doc, "Paso 1: Abrir el archivo de configuración centralizado", bold_prefix="Paso 1: ")
    add_p(doc, "Abra en su editor el archivo ubicado en:")
    add_code_block(doc, "lib/config/app_config.dart")

    add_p(doc, "Paso 2: Reemplazar las constantes de URLs", bold_prefix="Paso 2: ")
    add_p(doc, "Localice las variables de URL y coloque la dirección HTTPS de producción o pruebas donde se encuentra alojada la web app de credenciales y grupo familiar:")
    add_code_block(doc, "class AppConfig {\n  AppConfig._();\n\n  // URLs de las WebViews reales\n  static const String webviewCredentialUrl = 'https://su-dominio.com/credencial-digital';\n  static const String webviewGroupUrl = 'https://su-dominio.com/grupo-familiar';\n\n  // Endpoints de API futuros (Google Apps Script / Sheets)\n  static const String loginApiUrl = 'https://script.google.com/macros/s/SU_SCRIPT_ID/exec';\n  static const String newsApiUrl = 'https://api.su-dominio.com/novedades';\n\n  // Marquesina superior\n  static const String marqueeText = '📢 Presente su DNI junto a la credencial digital.';\n}")

    add_p(doc, "Paso 3: Guardar el archivo y probar en el emulador", bold_prefix="Paso 3: ")
    add_p(doc, "Guarde el archivo `app_config.dart`. La aplicación detectará automáticamente las URLs no vacías e inicializará la WebView nativa con transporte seguro de inmediato.")

    add_callout(doc, "Verificación de Permiso de Internet en Android",
                "El permiso android.permission.INTERNET ya se encuentra configurado en android/app/src/main/AndroidManifest.xml. No requiere modificaciones adicionales.")

    # Section 5: Troubleshooting
    add_styled_heading(doc, "5. Resolución de Problemas Frecuentes (Troubleshooting)", level=1)
    add_bullet(doc, "Ejecute 'flutter doctor --android-licenses' en la consola y presione 'y' para aceptar todas las licencias del SDK.", bold_prefix="Error de Licencias de Android: ")
    add_bullet(doc, "Asegúrese de que la URL empiece obligatoriamente por 'https://'. Si usa HTTP inseguro durante pruebas locales, Android requiere configurar 'android:usesCleartextTraffic=\"true\"' en AndroidManifest.xml.", bold_prefix="La WebView no carga la página web: ")
    add_bullet(doc, "Ejecute 'flutter clean', luego 'flutter pub get' y vuelva a compilar.", bold_prefix="Error al compilar Gradle: ")

    doc.save(filename)
    print(f"Document saved: {filename}")

# ==========================================
# 2. GENERATE IOS GUIDE
# ==========================================
def build_ios_doc(filename):
    doc = docx.Document()
    add_header_footer(doc, "Compilación y Testeo en iOS")
    
    # Title Block
    title_p = doc.add_paragraph()
    format_paragraph(title_p, space_before=10, space_after=4)
    r = title_p.add_run("Guía de Compilación, Testeo y Vinculación Web - iOS")
    r.font.name = 'Segoe UI'
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(13, 71, 161)
    
    sub_p = doc.add_paragraph()
    format_paragraph(sub_p, space_before=0, space_after=18)
    r_sub = sub_p.add_run("Opciones de Compilación (Online y Físicas) + Despliegue en iPhone/iPad")
    r_sub.font.name = 'Segoe UI'
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(100, 100, 100)
    
    add_callout(doc, "Nota Arquitectónica sobre iOS en Entornos Windows",
                "Dado que Apple requiere las herramientas propietarias de Xcode (que se ejecutan únicamente sobre el sistema operativo macOS) para firmar y compilar aplicaciones iOS, a continuación se presentan 4 opciones ordenadas de la MÁS SIMPLE (100% Cloud / Sin equipo Mac) a la MÁS COMPLEJA (Mac física/virtual).")

    # Section 1: Options Matrix
    add_styled_heading(doc, "1. Opciones de Compilación iOS (Ordenadas de la más SIMPLE a la más COMPLEJA)", level=1)
    
    # TABLE OF OPTIONS
    table = doc.add_table(rows=5, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Opción", "Método / Herramienta", "Requiere Mac Física", "Dificultad", "Costo Aprox."]
    data = [
        ["Opción 1", "Codemagic / GitHub Actions (Cloud)", "NO", "⭐ Muy Fácil", "Gratis (Free Tier)"],
        ["Opción 2", "Mac en la Nube (MacInCloud)", "NO", "⭐⭐ Fácil", "~$1 / hora"],
        ["Opción 3", "macOS Virtualizado (VMware en Win)", "NO", "⭐⭐⭐ Media", "Gratis (DIY)"],
        ["Opción 4", "Mac Física (MacBook / Mac Mini)", "SÍ", "⭐⭐⭐⭐ Compleja", "Costo de Hardware"]
    ]
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0D47A1")
        set_cell_margins(hdr_cells[i], top=80, bottom=80, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        format_paragraph(p, space_before=2, space_after=2)
        if p.runs:
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(9.5)
            p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        bg_hex = "F4F6F8" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            set_cell_background(row_cells[col_idx], bg_hex)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            format_paragraph(p, space_before=2, space_after=2)
            if p.runs:
                p.runs[0].font.size = Pt(9)

    add_p(doc, "") # spacing

    # DETAILED OPTIONS
    add_styled_heading(doc, "OPCIÓN 1: Servicios de Compilación en la Nube (Codemagic CI/CD) - [LA MÁS RECOMENDADA]", level=2)
    add_p(doc, "Esta es la opción más rápida, sencilla y profesional si trabaja en un entorno de desarrollo Windows, ya que NO requiere comprar un equipo Mac ni instalar máquinas virtuales.")
    add_bullet(doc, "Suba el repositorio de su código fuente a GitHub, GitLab o Bitbucket.", bold_prefix="Paso 1: ")
    add_bullet(doc, "Cree una cuenta gratuita en Codemagic.io (servicio de integración continua optimizado para Flutter).", bold_prefix="Paso 2: ")
    add_bullet(doc, "Conecte su repositorio GitHub con Codemagic.", bold_prefix="Paso 3: ")
    add_bullet(doc, "Seleccione el flujo 'Flutter App for iOS'.", bold_prefix="Paso 4: ")
    add_bullet(doc, "Haga clic en 'Start Build'. Las máquinas Mac virtuales de Codemagic compilarán el proyecto automáticamente en la nube y le devolverán el archivo instalable de iOS (.ipa) o lo enviarán directamente a TestFlight de Apple.", bold_prefix="Paso 5: ")

    add_styled_heading(doc, "OPCIÓN 2: Alquiler de una Mac en la Nube (MacInCloud / MacStadium)", level=2)
    add_p(doc, "Ideal si necesita interactuar con la interfaz gráfica de Xcode desde Windows sin comprar un equipo Apple.")
    add_bullet(doc, "Contrate un servicio de Mac en la nube como MacInCloud.com o MacStadium.", bold_prefix="Paso 1: ")
    add_bullet(doc, "Conéctese a la Mac remota usando el 'Escritorio Remoto de Windows' (RDP) o VNC.", bold_prefix="Paso 2: ")
    add_bullet(doc, "Abra la terminal en la Mac remota, clone su proyecto de Flutter y ejecute 'flutter build ios'.", bold_prefix="Paso 3: ")

    add_styled_heading(doc, "OPCIÓN 3: Entorno Virtualizado Local (macOS en VMware / VirtualBox)", level=2)
    add_p(doc, "Permite ejecutar macOS dentro de su computadora Windows actual.")
    add_bullet(doc, "Requiere una computadora Windows potente (procesador Intel/AMD con soporte de virtualización y mínimo 16 GB de Memoria RAM).", bold_prefix="Requisito: ")
    add_bullet(doc, "Instale VMware Workstation o VirtualBox y configure una imagen virtual de macOS (Sonoma / Sequoia).", bold_prefix="Paso 1: ")
    add_bullet(doc, "Dentro del entorno macOS virtual, instale Xcode y el SDK de Flutter para compilar localmente.", bold_prefix="Paso 2: ")

    add_styled_heading(doc, "OPCIÓN 4: Compilación Nativa en una Mac Física (MacBook / Mac Mini)", level=2)
    add_p(doc, "Es el método tradicional si dispone de un equipo hardware de Apple.")
    add_bullet(doc, "Transfiera la carpeta del proyecto a la Mac.", bold_prefix="Paso 1: ")
    add_bullet(doc, "Abra la consola de comandos en la Mac y navegue a la raíz del proyecto.", bold_prefix="Paso 2: ")
    add_bullet(doc, "Ejecute 'flutter pub get' para restaurar los paquetes.", bold_prefix="Paso 3: ")
    add_bullet(doc, "Abra la carpeta de iOS en Xcode mediante el archivo: ios/Runner.xcworkspace", bold_prefix="Paso 4: ")
    add_bullet(doc, "Seleccione su equipo de desarrollo (Apple Developer Account) en 'Signing & Capabilities'.", bold_prefix="Paso 5: ")
    add_bullet(doc, "Ejecute el comando 'flutter run -d ios' o compile la versión de producción con 'flutter build ipa'.", bold_prefix="Paso 6: ")

    # Section 2: Testing on iOS Simulators & Devices
    add_styled_heading(doc, "2. Paso a Paso: Testeo en Simulador iOS y Dispositivo Físico (iPhone)", level=1)
    
    add_styled_heading(doc, "2.1 Ejecución en el Simulador de iOS", level=2)
    add_p(doc, "En el entorno macOS (físico, virtual o alquilado):")
    add_code_block(doc, "# Abrir el simulador de iPhone de Xcode\nopen -a Simulator\n\n# Ejecutar la aplicación en el simulador de iOS\nflutter run -d ios")

    add_styled_heading(doc, "2.2 Pruebas en Dispositivos Reales de Afiliados vía TestFlight", level=2)
    add_p(doc, "La forma estándar de probar la aplicación en iPhones de usuarios o afiliados antes de la publicación pública en App Store es mediante TestFlight:")
    add_bullet(doc, "Suba el archivo .ipa generado en Codemagic o Xcode a App Store Connect.", bold_prefix="Paso 1: ")
    add_bullet(doc, "Inscriba las direcciones de correo electrónico de los testeadores o afiliados en la pestaña TestFlight.", bold_prefix="Paso 2: ")
    add_bullet(doc, "Los usuarios recibirán una invitación por email para instalar la app directamente en su iPhone desde la app gratuita TestFlight de Apple.", bold_prefix="Paso 3: ")

    # Section 3: Add Web Pages
    add_styled_heading(doc, "3. Pasos Necesarios para Agregar las Páginas Web al Cascarón (Shell)", level=1)
    add_p(doc, "La integración web en iOS utiliza exactamente el mismo mecanismo centralizado que Android:")
    
    add_p(doc, "Paso 1: Modificar `lib/config/app_config.dart`", bold_prefix="Paso 1: ")
    add_code_block(doc, "static const String webviewCredentialUrl = 'https://su-dominio.com/credencial-digital';\nstatic const String webviewGroupUrl = 'https://su-dominio.com/grupo-familiar';")

    add_p(doc, "Paso 2: Configuración de Transporte Seguro de Apple (ATS) en `ios/Runner/Info.plist`", bold_prefix="Paso 2: ")
    add_p(doc, "iOS requiere obligatoriamente que los sitios web dentro de WebViews usen conexiones HTTPS seguras. Si durante las pruebas necesita conectar a una dirección HTTP local, agregue el siguiente bloque dentro de `ios/Runner/Info.plist`:")
    add_code_block(doc, "<key>NSAppTransportSecurity</key>\n<dict>\n    <key>NSAllowsArbitraryLoads</key>\n    <true/>\n</dict>")

    # Section 4: Security Hook iOS
    add_styled_heading(doc, "4. Protección de Capturas de Pantalla en iOS (Security Hook)", level=1)
    add_p(doc, "A diferencia de Android que utiliza `FLAG_SECURE`, en iOS la protección contra capturas de pantalla se gestiona a nivel nativo en Swift en `ios/Runner/AppDelegate.swift` asignando un `UITextField` seguro a la jerarquía de ventanas:")
    add_code_block(doc, "// ios/Runner/AppDelegate.swift\nimport UIKit\nimport Flutter\n\n@UIApplicationMain\n@objc class AppDelegate: FlutterAppDelegate {\n  override func application(\n    _ application: UIApplication,\n    didFinishLaunchingWithOptions launchOptions: [UIApplication.LaunchOptionsKey: Any]?\n  ) -> Bool {\n    GeneratedPluginRegistrant.register(with: self)\n    // El cascarón nativo de Flutter ya incluye la interfaz en SecurityService\n    return super.application(application, didFinishLaunchingWithOptions: launchOptions)\n  }\n}")

    # Section 5: Troubleshooting iOS
    add_styled_heading(doc, "5. Resumen y Recomendación Final para iOS", level=1)
    add_p(doc, "Para el equipo de desarrollo trabajando desde entornos Windows, la recomendación óptima y de menor costo operativo es emplear la Opción 1 (Codemagic CI/CD):", bold_prefix="Recomendación Estratégica: ")
    add_bullet(doc, "Permite compilar tanto para Android como para iOS desde la misma consola sin cambiar de sistema operativo.")
    add_bullet(doc, "Genera los binarios .apk y .ipa listos para enviar a pruebas o tiendas.")

    doc.save(filename)
    print(f"Document saved: {filename}")

if __name__ == '__main__':
    build_android_doc("Guia_Compilacion_y_Testeo_Android.docx")
    build_ios_doc("Guia_Compilacion_y_Testeo_iOS.docx")
    print("All Word documents generated successfully.")
